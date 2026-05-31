"""Об'єднує дрібні абзаци в розділах 1-2 диплома Богдана.

Бере оригінал ``диплом 1_BACKUP_до_інтеграції.docx`` (бекап), знаходить
короткі абзаци що не є заголовками, формулами чи поясненнями до формул,
і склеює їх з попереднім сусідом. Результат — у файл
``диплом 1_BACKUP_merged.docx``, який потім використовує
assemble_full_diploma.py як теоретичну частину.

Запуск::

    python scripts/merge_small_paragraphs.py
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document


REPO_ROOT = Path(__file__).resolve().parent.parent
BACKUP_PATH = REPO_ROOT.parent / "диплом 1_BACKUP_до_інтеграції.docx"
OUT_PATH = REPO_ROOT.parent / "диплом 1_BACKUP_merged.docx"


# Регулярки для розпізнавання
RE_HEADING = re.compile(r"^\s*\d+\.\d+(\.\d+)?\s+")  # 1.1, 2.3.1
RE_CHAPTER = re.compile(r"^\s*(Висновки до розділу|РОЗДІЛ|Розділ)")
RE_FORMULA_VAR = re.compile(r"^\s*[А-Яа-яa-zA-ZA-Я]\s*[-—–]\s")  # "X - залежна змінна"
RE_LIST_DASH = re.compile(r"^\s*[-—–•]\s")
RE_FORMULA_NUM = re.compile(r"^\s*\(\d+\.\d+\)\s*$")  # (1.1)


def _is_heading(text: str) -> bool:
    return bool(RE_HEADING.match(text)) or bool(RE_CHAPTER.match(text))


def _is_formula_block(text: str) -> bool:
    """Рядок, що описує формулу або змінну формули."""
    t = text.strip()
    if not t:
        return True
    if t.startswith(("Де :", "Де:", "де :", "де:")):
        return True
    if RE_FORMULA_VAR.match(t):
        return True
    if RE_LIST_DASH.match(t):
        return True
    if RE_FORMULA_NUM.match(t):
        return True
    # рядки, що складаються переважно зі спецсимволів/математики
    if len(t) <= 4:
        return True
    return False


def _is_mergeable_prose(text: str) -> bool:
    """Звичайний прозовий абзац, який можна об'єднувати."""
    if not text or not text.strip():
        return False
    if _is_heading(text):
        return False
    if _is_formula_block(text):
        return False
    return True


def merge_in_place(doc, min_len: int = 250, max_combined: int = 1200) -> int:
    """Об'єднує сусідні малі прозові абзаци. Повертає кількість злиттів."""
    body = doc.element.body
    paragraphs = list(doc.paragraphs)
    n_merged = 0

    i = 1
    while i < len(paragraphs):
        cur = paragraphs[i]
        prev = paragraphs[i - 1]
        cur_text = cur.text.strip()
        prev_text = prev.text.strip()

        # умови злиття:
        #  - обидва — звичайні прозові абзаци
        #  - попередній занадто короткий АБО поточний дуже короткий
        #  - сумарна довжина не перевищує ліміт
        if (_is_mergeable_prose(prev_text) and _is_mergeable_prose(cur_text) and
                (len(prev_text) < min_len or len(cur_text) < 120) and
                len(prev_text) + len(cur_text) < max_combined):
            # перенести runs з cur у prev із пробілом-роздільником
            # спочатку — пробіл
            sep_run = prev.add_run(" ")
            # скопіювати базове форматування з останнього run попереднього
            if prev.runs:
                src = prev.runs[0]
                sep_run.font.name = src.font.name
                if src.font.size is not None:
                    sep_run.font.size = src.font.size
            # перенести усі runs з cur
            for run in cur.runs:
                new_run = prev.add_run(run.text)
                if run.font.name:
                    new_run.font.name = run.font.name
                if run.font.size is not None:
                    new_run.font.size = run.font.size
                new_run.font.bold = run.font.bold
                new_run.font.italic = run.font.italic
            # видалити cur з XML
            cur._element.getparent().remove(cur._element)
            # перебудувати список параграфів
            paragraphs = list(doc.paragraphs)
            n_merged += 1
            # i не інкрементуємо — продовжуємо аналізувати наступний
            # відносно вже об'єднаного prev
        else:
            i += 1

    return n_merged


def main() -> int:
    if not BACKUP_PATH.exists():
        print(f"ПОМИЛКА: {BACKUP_PATH} не знайдено", file=sys.stderr)
        return 1

    print(f"Відкриваю: {BACKUP_PATH.name}")
    doc = Document(str(BACKUP_PATH))

    before_total = len(doc.paragraphs)
    before_non_empty = sum(1 for p in doc.paragraphs if p.text.strip())
    before_short = sum(1 for p in doc.paragraphs
                       if 0 < len(p.text.strip()) < 250
                       and _is_mergeable_prose(p.text))
    print(f"  До: {before_total} параграфів усього, "
          f"{before_non_empty} непорожніх, {before_short} коротких прозових")

    n_merged = merge_in_place(doc)

    after_total = len(doc.paragraphs)
    after_non_empty = sum(1 for p in doc.paragraphs if p.text.strip())
    after_short = sum(1 for p in doc.paragraphs
                      if 0 < len(p.text.strip()) < 250
                      and _is_mergeable_prose(p.text))
    print(f"  Виконано злиттів: {n_merged}")
    print(f"  Після: {after_total} параграфів усього, "
          f"{after_non_empty} непорожніх, {after_short} коротких прозових")

    doc.save(str(OUT_PATH))
    kb = OUT_PATH.stat().st_size // 1024
    print(f"Збережено: {OUT_PATH.name} ({kb} KB)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
