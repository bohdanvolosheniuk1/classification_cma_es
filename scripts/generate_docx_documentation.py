"""Генерує формальну програмну документацію у форматі Word (.docx).

Цей документ є додатком до дипломної роботи Богдана. Структура
відповідає стандартним вимогам до програмної документації:
титульна сторінка, анотація, зміст, опис модулів, інструкція
користувача, опис алгоритмів, тестування, висновки, додатки.

Виклик::

    python scripts/generate_docx_documentation.py

Вихідний файл::

    ../docs/program_documentation.docx
"""

from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parent.parent
DOCS_DIR = REPO_ROOT.parent / "docs"
FIGURES_DIR = DOCS_DIR / "figures"
OUT_PATH = DOCS_DIR / "program_documentation.docx"


# ============================================================================
# допоміжні

def _set_run_style(run, *, size=12, bold=False, italic=False, font="Times New Roman",
                   color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    # для кирилиці треба прописати East Asia теж
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rfonts.set(qn("w:cs"), font)


def _para(doc, text, *, size=12, bold=False, italic=False, align=None,
          first_line_indent=Cm(1.25), space_after=Pt(6), line_spacing=1.5):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    pf.space_after = space_after
    pf.line_spacing = line_spacing
    run = p.add_run(text)
    _set_run_style(run, size=size, bold=bold, italic=italic)
    return p


def _heading(doc, text, level=1):
    sizes = {1: 16, 2: 14, 3: 13}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if level > 1 else WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(18 if level == 1 else 12)
    pf.space_after = Pt(12 if level == 1 else 8)
    pf.line_spacing = 1.15
    pf.keep_with_next = True
    run = p.add_run(text)
    _set_run_style(run, size=sizes.get(level, 12), bold=True)
    return p


def _bullet(doc, text, *, size=12):
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.space_after = Pt(3)
    pf.line_spacing = 1.5
    run = p.add_run(text)
    _set_run_style(run, size=size)
    return p


def _code(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(0.5)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    for i, line in enumerate(text.split("\n")):
        if i > 0:
            p.add_run().add_break()
        run = p.add_run(line)
        _set_run_style(run, size=10, font="Consolas",
                       color=RGBColor(0x33, 0x33, 0x33))
    return p


def _image(doc, path: Path, width_cm: float = 14, caption: str | None = None):
    if not path.exists():
        _para(doc, f"[не знайдено зображення: {path.name}]", italic=True,
              first_line_indent=None)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(12)
        run = cp.add_run(caption)
        _set_run_style(run, size=11, italic=True)


def _table(doc, headers: list[str], rows: list[list[str]],
           col_widths_cm: list[float] | None = None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        run = hdr_cells[i].paragraphs[0].add_run(h)
        _set_run_style(run, size=11, bold=True)
    for r_i, row in enumerate(rows, start=1):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i].cells[c_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            _set_run_style(run, size=11)
    if col_widths_cm is not None:
        for row in table.rows:
            for i, w in enumerate(col_widths_cm):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()  # spacing


def _page_break(doc):
    doc.add_page_break()


# ============================================================================
# контент

def _add_title_page(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("МІНІСТЕРСТВО ОСВІТИ І НАУКИ УКРАЇНИ")
    _set_run_style(run, size=14, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ЧЕРНІВЕЦЬКИЙ НАЦІОНАЛЬНИЙ УНІВЕРСИТЕТ\n"
                    "ІМЕНІ ЮРІЯ ФЕДЬКОВИЧА")
    _set_run_style(run, size=14, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Кафедра прикладної математики та інформаційних технологій")
    _set_run_style(run, size=12, italic=True)

    for _ in range(8):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ПРОГРАМНА ДОКУМЕНТАЦІЯ")
    _set_run_style(run, size=22, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("до дипломної роботи бакалавра")
    _set_run_style(run, size=14, italic=True)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("classification_cma_es:\nпорівняння класифікаторів із застосуванням\n"
                    "розширеного алгоритму CMA-ES")
    _set_run_style(run, size=16, bold=True)

    for _ in range(10):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Чернівці — {date.today().year}")
    _set_run_style(run, size=12)
    _page_break(doc)


def _add_annotation(doc):
    _heading(doc, "АНОТАЦІЯ", level=1)
    _para(doc,
        "Документ описує програмну реалізацію практичної частини дипломної "
        "роботи з прикладної математики. Програмний модуль "
        "classification_cma_es призначений для проведення порівняльного "
        "аналізу п'яти класифікаторів машинного навчання на трьох сучасних "
        "наборах даних із застосуванням розширеного алгоритму CMA-ES "
        "(Covariance Matrix Adaptation Evolution Strategy) зі сумішами "
        "нормальних розподілів, описаного в дисертації Літвінчук Ю.А. "
        "(Чернівецький національний університет, 2024)."
    )
    _para(doc,
        "Окремо реалізовано класифікатор на основі узагальнених адитивних "
        "моделей (GAM), теоретичні засади якого розглянуто в розділі 1 "
        "теоретичної частини диплома. Модель tuned_gam (GAM із "
        "гіперпараметрами, підібраними CMA-ES) являє собою практичну "
        "інтеграцію обох теоретичних розділів дипломної роботи."
    )
    _para(doc,
        "Програма побудована як модульний пакет мовою Python із "
        "застосуванням бібліотек scikit-learn, MLflow і Streamlit. "
        "Має командний і графічний інтерфейс. Покрита тестами pytest "
        "(34 тести)."
    )
    _para(doc,
        "Ключові слова: класифікація, CMA-ES, GAM, scikit-learn, MLflow, "
        "Streamlit, перехресна валідація, узагальнені адитивні моделі, "
        "еволюційні алгоритми, суміші розподілів.",
        italic=True
    )
    _page_break(doc)


def _add_toc(doc):
    _heading(doc, "ЗМІСТ", level=1)
    items = [
        ("ВСТУП", "5"),
        ("1. ЗАГАЛЬНІ ВІДОМОСТІ ПРО ПРОГРАМУ", "6"),
        ("    1.1. Назва і призначення", "6"),
        ("    1.2. Зв'язок з теоретичною частиною диплома", "6"),
        ("    1.3. Функціональні можливості", "7"),
        ("    1.4. Програмні засоби розробки", "7"),
        ("2. УМОВИ ВИКОРИСТАННЯ", "8"),
        ("    2.1. Апаратні вимоги", "8"),
        ("    2.2. Програмні вимоги", "8"),
        ("    2.3. Структура проекту", "9"),
        ("3. ВСТАНОВЛЕННЯ І НАЛАШТУВАННЯ", "10"),
        ("    3.1. Встановлення середовища", "10"),
        ("    3.2. Завантаження датасетів", "10"),
        ("    3.3. Перевірка установки", "11"),
        ("4. ОПИС МОДУЛІВ ПРОГРАМИ", "12"),
        ("    4.1. Модуль завантаження даних", "12"),
        ("    4.2. Модуль препроцесингу", "13"),
        ("    4.3. Базові класифікатори", "14"),
        ("    4.4. GAM-класифікатор", "15"),
        ("    4.5. Класичний CMA-ES", "16"),
        ("    4.6. Розширений CMA-ES зі сумішами", "17"),
        ("    4.7. CMA-NN", "19"),
        ("    4.8. Підбір гіперпараметрів", "20"),
        ("    4.9. Перехресна валідація", "21"),
        ("    4.10. Метрики", "22"),
        ("    4.11. MLflow-трекінг", "23"),
        ("    4.12. Pipeline", "24"),
        ("5. ІНТЕРФЕЙСИ КОРИСТУВАЧА", "25"),
        ("    5.1. Командний інтерфейс (CLI)", "25"),
        ("    5.2. Графічний інтерфейс (Streamlit)", "26"),
        ("    5.3. MLflow UI", "27"),
        ("6. ВХІДНІ ТА ВИХІДНІ ДАНІ", "28"),
        ("    6.1. Опис датасетів", "28"),
        ("    6.2. Формат вихідних метрик", "29"),
        ("7. ВИКОРИСТАНІ АЛГОРИТМИ", "30"),
        ("    7.1. Узагальнена адитивна модель (GAM)", "30"),
        ("    7.2. Класичний CMA-ES", "31"),
        ("    7.3. Розширений CMA-ES зі сумішами", "32"),
        ("    7.4. EM-алгоритм", "33"),
        ("8. ТЕСТУВАННЯ", "34"),
        ("ВИСНОВКИ", "35"),
        ("ДОДАТКИ", "36"),
        ("    А. Перелік моделей", "36"),
        ("    Б. Приклади CLI-команд", "37"),
        ("    В. Структура файлів", "38"),
    ]
    for title, page in items:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.tab_stops.add_tab_stop(Cm(16), WD_ALIGN_PARAGRAPH.RIGHT)
        pf.space_after = Pt(2)
        run = p.add_run(f"{title}\t{page}")
        _set_run_style(run, size=12)
    _page_break(doc)


def _add_intro(doc):
    _heading(doc, "ВСТУП", level=1)
    _para(doc,
        "Дипломна робота присвячена дослідженню узагальнених адитивних "
        "моделей (GAM) та алгоритму еволюційної стратегії з адаптацією "
        "коваріаційної матриці (CMA-ES). Практична частина роботи — програмний "
        "модуль classification_cma_es — реалізує обидва теоретичні напрямки "
        "у вигляді працюючого аналітичного інструменту для порівняння "
        "класифікаторів."
    )
    _para(doc,
        "Модуль розв'язує конкретну задачу куратора: провести порівняльний "
        "аналіз класифікації на трьох сучасних наборах даних із використанням "
        "п'яти різних методів класифікації, включно з алгоритмом, дослідженим "
        "у роботі Літвінчук Ю.А. (2024)."
    )
    _para(doc,
        "Документ структуровано наступним чином: розділ 1 містить загальні "
        "відомості про програму та її зв'язок з теорією диплома; розділ 2 "
        "описує умови використання та апаратні вимоги; розділ 3 — процедуру "
        "встановлення; розділ 4 — детальний опис кожного програмного "
        "модуля; розділ 5 — інтерфейси користувача; розділ 6 — опис "
        "вхідних і вихідних даних; розділ 7 — використані алгоритми; "
        "розділ 8 — тестування. У висновках підсумовано результати, "
        "у додатках наведено технічні таблиці."
    )
    _page_break(doc)


def _add_section_1(doc):
    _heading(doc, "1. ЗАГАЛЬНІ ВІДОМОСТІ ПРО ПРОГРАМУ", level=1)

    _heading(doc, "1.1. Назва і призначення", level=2)
    _para(doc, "Назва програмного засобу: classification_cma_es.")
    _para(doc,
        "Призначення: проведення порівняльного аналізу алгоритмів "
        "класифікації машинного навчання з можливістю автоматизованого "
        "підбору гіперпараметрів за допомогою еволюційної стратегії "
        "CMA-ES. Програма дозволяє завантажити вибраний датасет, "
        "виконати препроцесинг, провести розбиття на тренувальну і "
        "тестову вибірки, виконати k-fold перехресну валідацію, "
        "навчити кожну з обраних моделей, обчислити метрики якості "
        "(accuracy, F1-score, ROC-AUC) і порівняти результати у "
        "табличному та графічному вигляді."
    )

    _heading(doc, "1.2. Зв'язок з теоретичною частиною диплома", level=2)
    _para(doc,
        "Теоретична частина дипломної роботи містить два розділи: розділ 1 "
        "присвячено узагальненим адитивним моделям (GAM), розділ 2 — "
        "алгоритму CMA-ES. Програмна реалізація відображає обидва "
        "напрямки:"
    )
    _bullet(doc,
        "Розділ 1 теорії (GAM) реалізовано у класі GAMClassifier "
        "(модуль classifiers/gam.py). Використано B-сплайнове базисне "
        "перетворення кожної ознаки з наступною логістичною регресією, "
        "що відповідає GAM із логіт-функцією зв'язку (підрозділ 1.6 диплома)."
    )
    _bullet(doc,
        "Розділ 2 теорії (CMA-ES) реалізовано у двох варіантах: класичний "
        "(modul cma_es.py, обгортка над пакетом cma) та розширений зі "
        "сумішами нормальних розподілів (модуль mixture_cma_es.py, "
        "власна реалізація за дисертацією Літвінчук Ю.А., 2024)."
    )
    _bullet(doc,
        "Модель tuned_gam — це GAM, гіперпараметри якої (кількість "
        "вузлів сплайну k та параметр згладжування λ) підбираються "
        "автоматично за допомогою CMA-ES. Ця модель практично "
        "об'єднує обидва теоретичні розділи."
    )

    _heading(doc, "1.3. Функціональні можливості", level=2)
    _para(doc, "Програма забезпечує такі функції:")
    for item in [
        "автоматичне завантаження трьох сучасних датасетів "
        "(PhiUSIIL, Steel Plate Defects, Default of Credit Card Clients) "
        "з UCI Machine Learning Repository;",
        "опційне використання Kaggle-варіантів датасетів за наявності "
        "API-токену;",
        "уніфікований препроцесинг (імпутація пропусків, масштабування, "
        "one-hot енкодинг категоріальних ознак);",
        "розбиття на тренувальну/тестову вибірки зі стратифікацією;",
        "Stratified K-Fold перехресна валідація;",
        "навчання п'яти базових класифікаторів: логістична регресія, "
        "SVM, kNN, багатошарова нейронна мережа (MLP), GAM;",
        "навчання нейронної мережі (CMA-NN) за допомогою двох "
        "варіантів CMA-ES;",
        "автоматичний підбір гіперпараметрів кожної базової моделі "
        "за допомогою CMA-ES (моделі з префіксом tuned_);",
        "обчислення метрик: accuracy, F1-score (binary та weighted "
        "multiclass), ROC-AUC (binary та OvR weighted multiclass);",
        "журналювання експериментів у MLflow з можливістю "
        "відтворення;",
        "командний інтерфейс (CLI) для пакетних запусків;",
        "графічний інтерфейс (Streamlit) для інтерактивної роботи;",
        "експорт результатів у CSV-формат.",
    ]:
        _bullet(doc, item)

    _heading(doc, "1.4. Програмні засоби розробки", level=2)
    _table(doc,
        headers=["Засіб", "Версія", "Призначення"],
        rows=[
            ["Python", "3.10+", "Мова реалізації"],
            ["NumPy", "≥1.26", "Лінійна алгебра, тензорні обчислення"],
            ["pandas", "≥2.1", "Робота з табличними даними"],
            ["scikit-learn", "≥1.4", "Базові класифікатори, препроцесинг"],
            ["cma", "≥3.3", "Класичний CMA-ES"],
            ["mlflow", "≥2.10", "Трекінг експериментів"],
            ["streamlit", "≥1.30", "Графічний інтерфейс"],
            ["altair", "≥5.0", "Візуалізація"],
            ["ucimlrepo", "≥0.0.7", "Завантаження датасетів з UCI"],
            ["pytest", "≥8.0", "Модульне тестування"],
        ],
        col_widths_cm=[5, 3, 8])
    _page_break(doc)


def _add_section_2(doc):
    _heading(doc, "2. УМОВИ ВИКОРИСТАННЯ", level=1)

    _heading(doc, "2.1. Апаратні вимоги", level=2)
    _table(doc,
        headers=["Компонент", "Мінімум", "Рекомендовано"],
        rows=[
            ["Процесор", "x86-64, 2 ядра, 2 ГГц", "x86-64, 4+ ядра, 3+ ГГц"],
            ["Оперативна память", "4 ГБ", "8+ ГБ"],
            ["Дисковий простір", "2 ГБ вільно", "5 ГБ вільно"],
            ["Підключення до мережі", "потрібно для перш. завантаження датасетів", "—"],
        ],
        col_widths_cm=[5, 5, 6])

    _heading(doc, "2.2. Програмні вимоги", level=2)
    _para(doc,
        "Програма розроблена для запуску на операційних системах "
        "Windows 10/11, Linux або macOS. Усі залежності кросплатформові."
    )
    _para(doc, "Необхідне системне програмне забезпечення:")
    _bullet(doc, "Python 3.10 або вище (тестовано на 3.14.3).")
    _bullet(doc, "pip — менеджер пакетів Python.")
    _bullet(doc,
        "Git — система контролю версій (опційно, для клонування).")
    _bullet(doc,
        "Веб-браузер (Chrome, Firefox, Edge) для роботи з Streamlit "
        "та MLflow UI.")

    _heading(doc, "2.3. Структура проекту", level=2)
    _para(doc,
        "Програмний модуль організовано у вигляді стандартного Python-пакета:"
    )
    _code(doc,
        "classification_cma_es/           корінь проекту\n"
        "  classifiers/                   основний пакет\n"
        "    __init__.py\n"
        "    data.py                      завантаження датасетів\n"
        "    preprocessing.py             препроцесинг\n"
        "    models.py                    фабрики класифікаторів\n"
        "    gam.py                       GAM (розділ 1)\n"
        "    cma_es.py                    класичний CMA-ES\n"
        "    mixture_cma_es.py            розширений CMA-ES + EM\n"
        "    cma_nn.py                    нейромережа на CMA-ES\n"
        "    hyperparam_tuning.py         tuned_* моделі\n"
        "    crossval.py                  k-fold CV\n"
        "    metrics.py                   accuracy, F1, AUC\n"
        "    tracking.py                  MLflow\n"
        "    pipeline.py                  оркестратор\n"
        "  scripts/                       службові скрипти\n"
        "    run_experiment.py            CLI\n"
        "    download.py                  завантаження датасетів\n"
        "    generate_guide.py            PDF-гайд\n"
        "    generate_api_docs.py         HTML API\n"
        "    generate_docx_documentation.py  цей документ\n"
        "  tests/                         pytest-тести (34 шт.)\n"
        "  app.py                         Streamlit-дашборд\n"
        "  run_app.bat                    лаунчер дашборду\n"
        "  run_mlflow.bat                 лаунчер MLflow UI\n"
        "  setup.py                       конфігурація пакета\n"
        "  requirements.txt               список залежностей\n"
        "  README.md                      короткий опис\n"
        "  CHANGELOG.md                   історія змін\n"
        "  CONTRIBUTING.md                інструкція для розробників\n"
        "  LICENSE                        MIT"
    )
    _page_break(doc)


def _add_section_3(doc):
    _heading(doc, "3. ВСТАНОВЛЕННЯ І НАЛАШТУВАННЯ", level=1)

    _heading(doc, "3.1. Встановлення середовища", level=2)
    _para(doc, "Створення віртуального середовища і встановлення залежностей:")
    _code(doc,
        "# Windows PowerShell\n"
        "cd classification_cma_es\n"
        "py -3 -m venv .venv\n"
        ".\\.venv\\Scripts\\Activate.ps1\n"
        "pip install -e ."
    )
    _para(doc,
        "Команда pip install -e . встановить пакет classification_cma_es "
        "у редагованому режимі (editable install) разом з усіма залежностями, "
        "переліченими у файлі requirements.txt."
    )

    _heading(doc, "3.2. Завантаження датасетів", level=2)
    _para(doc,
        "Датасет PhiUSIIL завантажується автоматично при першому виклику "
        "load_dataset('phiusiil') через бібліотеку ucimlrepo. Для двох "
        "інших датасетів передбачено два варіанти:"
    )
    _bullet(doc,
        "Якщо встановлено пакет kaggle та надано API-токен (kaggle.json у "
        "~/.kaggle/) — використовуються Kaggle-варіанти (2024).")
    _bullet(doc,
        "Якщо Kaggle недоступний — автоматично активується UCI fallback: "
        "Steel Plate Faults (UCI #198) та Default of Credit Card Clients "
        "(UCI #350).")
    _para(doc, "Команда для пакетного завантаження:")
    _code(doc,
        "python scripts/download.py --dataset phiusiil\n"
        "python scripts/download.py --dataset steel_plate\n"
        "python scripts/download.py --dataset loan_approval"
    )

    _heading(doc, "3.3. Перевірка установки", level=2)
    _para(doc, "Запуск тестового набору:")
    _code(doc, "pytest -q")
    _para(doc,
        "Очікуваний результат: 34 passed. Усі модулі покриті тестами, "
        "включаючи завантаження датасетів, препроцесинг, кожен класифікатор, "
        "класичний і розширений CMA-ES, EM-алгоритм, метрики."
    )
    _page_break(doc)


def _add_section_4(doc):
    _heading(doc, "4. ОПИС МОДУЛІВ ПРОГРАМИ", level=1)

    # 4.1 data
    _heading(doc, "4.1. Модуль завантаження даних (data.py)", level=2)
    _para(doc,
        "Модуль надає уніфікований інтерфейс завантаження трьох датасетів "
        "у вигляді датакласу Dataset, що містить матрицю ознак X, цільову "
        "змінну y, тип задачі (binary/multiclass) і кількість класів."
    )
    _para(doc, "Основні публічні функції:")
    _bullet(doc, "load_dataset(name) — універсальний завантажувач за іменем.")
    _bullet(doc, "load_phiusiil() — PhiUSIIL Phishing URL (UCI #967, 2024).")
    _bullet(doc, "load_steel_plate() — Steel Plate Defects (UCI #198 fallback).")
    _bullet(doc, "load_loan_approval() — Default of Credit Card Clients (UCI #350).")
    _para(doc,
        "Усі функції підтримують локальний кеш у вигляді CSV-файлу — "
        "повторні виклики читають з диска без повторного звертання до UCI."
    )

    # 4.2 preprocessing
    _heading(doc, "4.2. Модуль препроцесингу (preprocessing.py)", level=2)
    _para(doc,
        "Реалізує побудову sklearn ColumnTransformer, який однаково "
        "обробляє будь-який із трьох датасетів. Числові ознаки проходять "
        "імпутацію медіаною і StandardScaler; категоріальні — найчастішим "
        "значенням і OneHotEncoder."
    )
    _para(doc, "Функція encode_target(y) кодує цільову змінну в цілі числа 0..K-1 "
              "із детермінованим порядком класів (sorted).")

    # 4.3 models
    _heading(doc, "4.3. Базові класифікатори (models.py)", level=2)
    _para(doc, "Модуль містить п'ять фабрик класифікаторів:")
    _table(doc,
        headers=["Фабрика", "Клас sklearn", "Дефолтні параметри"],
        rows=[
            ["make_logreg", "LogisticRegression", "max_iter=1000"],
            ["make_svm", "SVC", "C=1.0, kernel=rbf, probability=True"],
            ["make_knn", "KNeighborsClassifier", "n_neighbors=5"],
            ["make_mlp", "MLPClassifier", "hidden=(64, 32), max_iter=300"],
            ["make_gam", "GAMClassifier (власний)", "n_knots=5, degree=3, C=1.0"],
        ],
        col_widths_cm=[4, 5, 7])
    _para(doc,
        "Усі моделі сумісні зі sklearn API: метод fit(X, y) для навчання, "
        "predict(X) для прогнозів класів, predict_proba(X) для ймовірностей."
    )

    # 4.4 gam
    _heading(doc, "4.4. GAM-класифікатор (gam.py)", level=2)
    _para(doc,
        "Узагальнена адитивна модель — це власна реалізація поверх "
        "scikit-learn без сторонніх залежностей (pygam відсутній у "
        "PyPI для Python 3.14). Модель будується як sklearn Pipeline: "
        "SplineTransformer (B-сплайнове базисне перетворення кожної "
        "ознаки) → LogisticRegression. Математично еквівалентна GAM з "
        "логіт-функцією зв'язку."
    )
    _image(doc, FIGURES_DIR / "04_gam.png", width_cm=15,
           caption="Рисунок 4.1 — GAM як сума гладких функцій від ознак")
    _para(doc, "Параметри GAMClassifier:")
    _bullet(doc, "n_knots (за замовчуванням 5) — кількість вузлів сплайну k.")
    _bullet(doc, "degree (3) — степінь B-сплайну (3 = кубічний).")
    _bullet(doc, "C (1.0) — обернений параметр згладжування λ = 1/C.")
    _bullet(doc, "knots ('uniform') — стратегія розміщення вузлів.")

    # 4.5 cma_es
    _heading(doc, "4.5. Класичний CMA-ES (cma_es.py)", level=2)
    _para(doc,
        "Тонка обгортка над пакетом cma Ніколаса Хансена (той самий пакет "
        "використовує Літвінчук у дисертації). Функція minimize_cma "
        "приймає цільову функцію, початкову точку, початковий σ, "
        "опційні межі та повертає результат із полями best_x, best_f, "
        "n_evaluations, n_iterations, history."
    )
    _image(doc, FIGURES_DIR / "05_cmaes_cycle.png", width_cm=12,
           caption="Рисунок 4.2 — Цикл класичного CMA-ES")

    # 4.6 mixture
    _heading(doc, "4.6. Розширений CMA-ES зі сумішами (mixture_cma_es.py)", level=2)
    _para(doc,
        "Власна реалізація алгоритму, описаного в дисертації Літвінчук Ю.А. "
        "(2024). Унімодальний нормальний розподіл класичного CMA-ES "
        "замінено на суміш k нормальних компонент, параметри якої "
        "оцінюються EM-алгоритмом за найкращими хромосомами кожної "
        "ітерації."
    )
    _image(doc, FIGURES_DIR / "06_mixture.png", width_cm=15,
           caption="Рисунок 4.3 — Суміш 3-х нормальних і EM-цикл")
    _para(doc, "Клас MixtureCMAES має такі ключові параметри:")
    _bullet(doc, "n_components (k) — початкова кількість піків у суміші.")
    _bullet(doc, "pop_size (N) — кількість хромосом на ітерацію.")
    _bullet(doc, "em_steps — кількість EM-ітерацій на одне зовнішнє оновлення.")
    _bullet(doc,
        "cov_lr — швидкість оновлення коваріаційних матриць. "
        "Чисте EM (lr=1) призводить до передчасного колапсу дисперсії "
        "на унімодальних задачах, тому додано момент-усереднення "
        "коваріації (lr=0.2 за замовчуванням) як rank-μ-аналог "
        "класичного CMA-ES.")
    _bullet(doc,
        "adaptive (True/False) — самоадаптивний підбір k: видалення "
        "малих піків (|X_l| < √(N/2)) і додавання нового піку у разі "
        "стагнації цільової функції.")

    # 4.7 cma_nn
    _heading(doc, "4.7. CMA-NN — нейронна мережа з CMA-ES-навчанням (cma_nn.py)", level=2)
    _para(doc,
        "Клас CMAESNeuralNet — це власна реалізація невеликої повнозв'язної "
        "нейронної мережі, ваги якої навчаються не методом зворотного "
        "поширення (backpropagation), а оптимізатором CMA-ES. Усі ваги "
        "укладено в один плоский вектор w ∈ R^n; CMA-ES шукає таке w, "
        "яке мінімізує крос-ентропію на тренувальних даних. Цей "
        "класифікатор являє собою практичну реалізацію методу, "
        "досліджуваного у роботі Літвінчук Ю.А., — це і є той 5-й "
        "класифікатор, що зазначений у завданні куратора."
    )
    _image(doc, FIGURES_DIR / "07_cma_nn.png", width_cm=15,
           caption="Рисунок 4.4 — CMA-NN: ваги мережі як вектор для оптимізатора")
    _para(doc, "Підтримує два режими через параметр method:")
    _bullet(doc, "'classic' — використовує minimize_cma з cma_es.py.")
    _bullet(doc, "'mixture' — використовує MixtureCMAES зі сумішами.")
    _para(doc,
        "Для прискорення на широких ознаках передбачено опційний PCA "
        "(max_features) і субсемплування train (max_train_samples)."
    )

    # 4.8 hyperparam
    _heading(doc, "4.8. Підбір гіперпараметрів (hyperparam_tuning.py)", level=2)
    _para(doc,
        "Модуль реалізує моделі з префіксом tuned_, в яких CMA-ES "
        "використовується не для навчання моделі, а для пошуку її "
        "оптимальних гіперпараметрів. Для кожної базової моделі "
        "визначено простір HyperSpace із межами і функцією трансформації "
        "вектора CMA-ES у словник параметрів sklearn-моделі."
    )
    _para(doc, "Простори гіперпараметрів:")
    _table(doc,
        headers=["Модель", "Простір", "Розмірність"],
        rows=[
            ["tuned_logreg", "log10(C) ∈ [-3, 3]", "1"],
            ["tuned_svm", "log10(C) ∈ [-3, 3], log10(γ) ∈ [-4, 1]", "2"],
            ["tuned_knn", "n_neighbors ∈ [1, 30]", "1"],
            ["tuned_mlp", "hidden_size, log10(α), log10(lr)", "3"],
            ["tuned_gam", "n_knots, degree, log10(C)", "3"],
        ],
        col_widths_cm=[3.5, 8, 3])
    _para(doc,
        "Модель tuned_gam є особливою — це безпосередня практична "
        "інтеграція обох теоретичних розділів дипломної роботи: "
        "CMA-ES (розділ 2) шукає оптимальні значення параметрів GAM "
        "(розділ 1)."
    )

    # 4.9 crossval
    _heading(doc, "4.9. Перехресна валідація (crossval.py)", level=2)
    _para(doc, "Модуль інкапсулює дві стандартні процедури оцінки:")
    _bullet(doc,
        "split_train_test(X, y) — розбиття у пропорції 80/20 зі "
        "стратифікацією (зберігаються пропорції класів у вибірках).")
    _bullet(doc,
        "kfold_evaluate(factory, X, y, n_splits) — Stratified K-Fold "
        "із значенням k=5 за замовчуванням. На кожному фолді створюється "
        "свіжа модель і оцінюється на валідаційному підмножині.")
    _image(doc, FIGURES_DIR / "03_kfold.png", width_cm=15,
           caption="Рисунок 4.5 — Stratified K-Fold з k=5")

    # 4.10 metrics
    _heading(doc, "4.10. Метрики (metrics.py)", level=2)
    _para(doc, "Реалізує три метрики, затверджені куратором:")
    _bullet(doc,
        "Accuracy — частка правильних прогнозів. Не використовується як "
        "основна метрика на незбалансованих задачах.")
    _bullet(doc,
        "F1-score — для бінарних задач використовується стандартний F1 "
        "позитивного класу; для мультикласу — weighted average по класах.")
    _bullet(doc,
        "ROC-AUC — для бінарних задач звичайний ROC-AUC; для мультикласу — "
        "One-vs-Rest weighted average.")
    _para(doc,
        "Функція aggregate_folds(folds) обчислює середні значення (mean) "
        "і стандартні відхилення (std) метрик по фолдах для оцінки "
        "стабільності моделі."
    )

    # 4.11 tracking
    _heading(doc, "4.11. MLflow-трекінг (tracking.py)", level=2)
    _para(doc,
        "Модуль надає тонкий шар над MLflow API для автоматичного "
        "логування експериментів. Кожен запуск моделі реєструється як "
        "окремий MLflow Run у спільному експерименті. Логуються: "
        "параметри (датасет, модель, фолди, seed, кількість прикладів), "
        "метрики (test_*, cv_*, fit_time), для tuned_*-моделей — "
        "знайдені оптимальні гіперпараметри."
    )
    _para(doc,
        "Дані зберігаються локально у теці mlruns/ і доступні для "
        "перегляду через MLflow UI."
    )

    # 4.12 pipeline
    _heading(doc, "4.12. Pipeline (pipeline.py)", level=2)
    _para(doc,
        "Оркестратор експерименту — центральний модуль, що поєднує всі "
        "інші. Функція run_experiment приймає назву датасету, список "
        "моделей і параметри запуску, виконує повний цикл: завантаження "
        "→ препроцесинг → train/test split → для кожної моделі (k-fold "
        "CV → фінальна оцінка на test) → опційне MLflow-логування."
    )
    _image(doc, FIGURES_DIR / "02_pipeline.png", width_cm=16,
           caption="Рисунок 4.6 — Загальна послідовність експерименту")
    _para(doc,
        "Функція підтримує колбеки для прогресу (on_dataset_ready, "
        "on_model_start, on_model_done, on_model_error), які дозволяють "
        "відображати стан виконання у графічному інтерфейсі Streamlit."
    )
    _page_break(doc)


def _add_section_5(doc):
    _heading(doc, "5. ІНТЕРФЕЙСИ КОРИСТУВАЧА", level=1)

    _heading(doc, "5.1. Командний інтерфейс (CLI)", level=2)
    _para(doc, "Запуск експерименту з командного рядка:")
    _code(doc,
        "python scripts/run_experiment.py --dataset phiusiil \\\n"
        "    --models all --folds 5 --sample 5000"
    )
    _para(doc, "Доступні опції CLI:")
    _table(doc,
        headers=["Опція", "Опис"],
        rows=[
            ["--dataset", "phiusiil | steel_plate | loan_approval"],
            ["--models", "all або csv (logreg,svm,gam,...)"],
            ["--folds", "кількість фолдів k-fold CV (3..10)"],
            ["--sample N", "обмеження розміру вибірки (для прискорення)"],
            ["--cma-iter N", "ліміт ітерацій CMA-ES (15..120)"],
            ["--seed", "глобальний seed для відтворюваності"],
            ["--no-mlflow", "не логувати в MLflow"],
            ["--no-tuning", "пропустити моделі tuned_*"],
            ["--mlflow-uri", "URI MLflow-сервера (за замовчуванням локальний)"],
        ],
        col_widths_cm=[4.5, 11])

    _heading(doc, "5.2. Графічний інтерфейс (Streamlit)", level=2)
    _para(doc,
        "Графічний інтерфейс реалізовано на основі бібліотеки Streamlit. "
        "Запуск — подвійним кліком на файл run_app.bat у корені проекту "
        "(Windows), що автоматично активує віртуальне середовище і "
        "запускає Streamlit. Браузер відкривається автоматично на "
        "адресі http://localhost:8501."
    )
    _para(doc, "Структура дашборду:")
    _bullet(doc,
        "Бокова панель містить елементи керування: вибір датасету "
        "(dropdown), список моделей (multiselect), повзунок кількості "
        "фолдів, поле обмеження розміру вибірки, повзунок ітерацій "
        "CMA-ES, перемикач MLflow-логування, поле seed.")
    _bullet(doc,
        "Кнопка 'Запустити' розпочинає експеримент із відображенням "
        "прогресу у вигляді прогрес-бара і логу виконання кожної моделі.")
    _bullet(doc,
        "Після завершення відображаються: 5 карток зі статистикою "
        "(розмір train/test, кількість ознак і класів), таблиця "
        "результатів із кольоровими ProgressColumn для метрик "
        "(адаптивний zoom по фактичним значенням), кнопка експорту "
        "у CSV, гістограма часу навчання, крива збіжності для "
        "CMA-моделей.")
    _para(doc,
        "Стан останнього запуску зберігається на диск (results/_last_run.json) "
        "— у разі розриву websocket-зʼєднання дашборд відновлює "
        "результати при оновленні сторінки."
    )

    _heading(doc, "5.3. MLflow UI", level=2)
    _para(doc,
        "Перегляд журналу експериментів — подвійний клік на "
        "run_mlflow.bat. Відкривається http://localhost:5000 із "
        "повним списком запусків, можливістю фільтрації, сортування, "
        "експорту і побудови графіків."
    )
    _page_break(doc)


def _add_section_6(doc):
    _heading(doc, "6. ВХІДНІ ТА ВИХІДНІ ДАНІ", level=1)

    _heading(doc, "6.1. Опис датасетів", level=2)
    _table(doc,
        headers=["Датасет", "Тип задачі", "Розмір (UCI fallback)", "Рік"],
        rows=[
            ["PhiUSIIL", "бінарна", "235 795 × 54", "2024"],
            ["Steel Plate Defects", "мультикласова (7)", "1 941 × 27", "2010"],
            ["Default of Credit Card Clients", "бінарна", "30 000 × 23", "2016"],
        ],
        col_widths_cm=[5, 4, 4, 2.5])
    _para(doc,
        "PhiUSIIL — задача класифікації URL-адрес на phishing та "
        "легітимні. Ознаки: довжина URL, кількість крапок, наявність "
        "HTTPS, ентропія, інші структурні характеристики."
    )
    _para(doc,
        "Steel Plate Defects — задача класифікації типів дефектів "
        "металевих пластин: Pastry, Z_Scratch, K_Scratch, Stains, "
        "Dirtiness, Bumps, Other_Faults."
    )
    _para(doc,
        "Default of Credit Card Clients — задача прогнозування "
        "ймовірності дефолту клієнта банку у наступному місяці на "
        "основі кредитної історії, демографічних даних, "
        "статусу платежів. Класи незбалансовані (~78% / 22%)."
    )

    _heading(doc, "6.2. Формат вихідних метрик", level=2)
    _para(doc, "Для кожної моделі обчислюються:")
    _bullet(doc, "test_accuracy, test_f1, test_auc — на тестовій вибірці.")
    _bullet(doc,
        "cv_accuracy_mean, cv_f1_mean, cv_auc_mean — середні по фолдах "
        "k-fold CV.")
    _bullet(doc, "cv_accuracy_std, cv_f1_std, cv_auc_std — стандартні відхилення.")
    _bullet(doc, "fit_time_s — час навчання фінальної моделі у секундах.")
    _bullet(doc, "tune_time_s — додатково для tuned_* моделей — час підбору.")
    _para(doc,
        "Результати зберігаються у вигляді CSV-файлу "
        "results/<dataset>/summary.csv та автоматично логуються в MLflow."
    )
    _page_break(doc)


def _add_section_7(doc):
    _heading(doc, "7. ВИКОРИСТАНІ АЛГОРИТМИ", level=1)

    _heading(doc, "7.1. Узагальнена адитивна модель (GAM)", level=2)
    _para(doc,
        "Загальна форма GAM з функцією зв'язку g:"
    )
    _code(doc, "g(E[y|X]) = β₀ + Σᵢ fᵢ(xᵢ)")
    _para(doc,
        "де кожна fᵢ — гладка функція, що представляється у вигляді "
        "лінійної комбінації базисних функцій (зазвичай — B-сплайнів):"
    )
    _code(doc, "fᵢ(xᵢ) = Σⱼ βᵢⱼ · bⱼ(xᵢ)")
    _para(doc,
        "У нашій реалізації використано логіт-функцію зв'язку "
        "(класифікаційна задача), а коефіцієнти βᵢⱼ підбирає логістична "
        "регресія. Штраф L2 LogisticRegression виконує роль параметра "
        "згладжування λ."
    )

    _heading(doc, "7.2. Класичний CMA-ES", level=2)
    _para(doc,
        "Алгоритм еволюційної стратегії з адаптацією коваріаційної "
        "матриці. На кожній ітерації:"
    )
    _bullet(doc,
        "1. Семплування λ кандидатних точок із багатовимірного "
        "нормального розподілу xᵢ ~ N(m, σ²C).")
    _bullet(doc, "2. Обчислення цільової функції f(xᵢ) для кожного кандидата.")
    _bullet(doc, "3. Сортування за якістю, відбір μ найкращих.")
    _bullet(doc,
        "4. Оновлення вектора середніх m, коваріаційної матриці C та "
        "розміру кроку σ за відібраними кандидатами та шляхами еволюції.")

    _heading(doc, "7.3. Розширений CMA-ES зі сумішами", level=2)
    _para(doc,
        "Метод, описаний у дисертації Літвінчук Ю.А. (2024). "
        "Унімодальний нормальний розподіл замінено на суміш k "
        "нормальних компонент:"
    )
    _code(doc, "p(x) = Σₛ wₛ · N(x; mₛ, Cₛ), Σ wₛ = 1")
    _para(doc,
        "Параметри суміші (wₛ, mₛ, Cₛ) оновлюються EM-алгоритмом за "
        "найкращою половиною хромосом. Опційно — самоадаптивний підбір "
        "k: видалення малих піків і додавання нового у разі стагнації."
    )
    _para(doc,
        "Технічне доповнення (не з оригінальної дисертації): для "
        "запобігання передчасному колапсу дисперсії додано параметр "
        "cov_lr — момент-усереднення коваріації між ітераціями. Без цієї "
        "поправки алгоритм не сходиться навіть на простих унімодальних "
        "задачах (тестовано на функції сфери)."
    )

    _heading(doc, "7.4. EM-алгоритм для суміші нормальних", level=2)
    _para(doc, "E-крок (обчислення апостеріорних відповідальностей):")
    _code(doc, "γᵢⱼ = (wⱼ · N(xᵢ; mⱼ, Cⱼ)) / Σₛ (wₛ · N(xᵢ; mₛ, Cₛ))")
    _para(doc, "M-крок (оновлення параметрів суміші):")
    _code(doc,
        "Nⱼ  = Σᵢ γᵢⱼ\n"
        "wⱼ  = Nⱼ / N\n"
        "mⱼ  = (Σᵢ γᵢⱼ · xᵢ) / Nⱼ\n"
        "Cⱼ  = (Σᵢ γᵢⱼ · (xᵢ - mⱼ)(xᵢ - mⱼ)ᵀ) / Nⱼ"
    )
    _para(doc,
        "Обчислення проводяться у логарифмічній шкалі через log-sum-exp "
        "для чисельної стабільності. Коваріаційні матриці регуляризуються "
        "додаванням ε·I для уникнення виродженості."
    )
    _page_break(doc)


def _add_section_8(doc):
    _heading(doc, "8. ТЕСТУВАННЯ", level=1)
    _para(doc,
        "Програма покрита автоматизованими тестами у фреймворку pytest. "
        "Тестовий набір складається з 34 тестів, розподілених між 7 "
        "тестовими модулями:"
    )
    _table(doc,
        headers=["Модуль", "Кількість тестів", "Що перевіряє"],
        rows=[
            ["test_metrics.py", "4", "compute_metrics, aggregate_folds"],
            ["test_preprocessing.py", "3", "ColumnTransformer, encode_target"],
            ["test_cma_es.py", "3", "Сфера, Розенброк, межі для класичного CMA-ES"],
            ["test_mixture_cma.py", "5", "EM-крок, збіжність на сфері та бімодальній функції, адаптація k"],
            ["test_cma_nn.py", "4", "fit/predict для classic і mixture, PCA, субсемплування"],
            ["test_hyperparam.py", "6", "Простори, трансформації, tune_with_cma"],
            ["test_gam.py", "9", "GAM на binary і multiclass, різні n_knots, фабрика, простір, tuning"],
        ],
        col_widths_cm=[4, 2.5, 9])
    _para(doc, "Запуск повного набору:")
    _code(doc, "pytest -q")
    _para(doc, "Очікуваний результат — 34 passed, час виконання ≈ 7 секунд.")
    _page_break(doc)


def _add_conclusions(doc):
    _heading(doc, "ВИСНОВКИ", level=1)
    _para(doc,
        "Розроблено програмний модуль classification_cma_es, що реалізує "
        "практичну частину дипломної роботи з прикладної математики. Модуль "
        "повністю відповідає вимогам куратора:"
    )
    _bullet(doc,
        "Реалізовано п'ять класифікаторів машинного навчання, включно з "
        "методом, дослідженим у роботі Літвінчук Ю.А.")
    _bullet(doc,
        "Виконується стандартна методика оцінки: розбиття на тренувальну/"
        "тестову вибірки і k-fold перехресна валідація.")
    _bullet(doc,
        "Обчислюються три затверджені метрики: accuracy, F1-score, AUC.")
    _bullet(doc,
        "Програма організована як модульний Python-пакет із детальною "
        "документацією (NumPy-style docstring'и, HTML API через pdoc, "
        "PDF-гайд зі схемами, цей Word-документ).")
    _bullet(doc,
        "Реалізовано трекінг експериментів через MLflow з можливістю "
        "відтворення та порівняння запусків.")
    _para(doc, "Додатково, понад вимоги куратора, реалізовано:")
    _bullet(doc,
        "GAM-класифікатор, що відповідає теорії розділу 1 диплома, та "
        "модель tuned_gam, що інтегрує обидва теоретичні розділи.")
    _bullet(doc,
        "Розширений CMA-ES зі сумішами нормальних розподілів — власна "
        "реалізація за дисертацією Літвінчук Ю.А. з технічним "
        "доповненням (параметр cov_lr) для практичної стабільності.")
    _bullet(doc,
        "Графічний дашборд на Streamlit з адаптивними візуалізаціями.")
    _bullet(doc,
        "Лаунчери для запуску одним кліком (без термінала).")
    _bullet(doc,
        "Автоматизований тестовий набір (34 тести pytest).")
    _para(doc,
        "Модуль готовий до подальшого розширення (додавання нових "
        "датасетів або моделей) — процедура описана у файлі CONTRIBUTING.md."
    )
    _page_break(doc)


def _add_appendices(doc):
    _heading(doc, "ДОДАТКИ", level=1)

    _heading(doc, "Додаток А. Перелік моделей", level=2)
    _table(doc,
        headers=["Ім'я моделі", "Опис", "Категорія"],
        rows=[
            ["logreg", "Логістична регресія", "базова"],
            ["svm", "SVM з RBF-ядром", "базова"],
            ["knn", "k-найближчих сусідів", "базова"],
            ["mlp", "MLPClassifier sklearn", "базова"],
            ["gam", "GAM-класифікатор (розділ 1)", "базова"],
            ["cma_classic", "NN, навчена класичним CMA-ES", "5-й метод"],
            ["cma_mixture", "NN, навчена розширеним CMA-ES", "5-й метод"],
            ["tuned_logreg", "logreg + CMA-ES tuning", "tuned"],
            ["tuned_svm", "svm + CMA-ES tuning", "tuned"],
            ["tuned_knn", "knn + CMA-ES tuning", "tuned"],
            ["tuned_mlp", "mlp + CMA-ES tuning", "tuned"],
            ["tuned_gam", "gam + CMA-ES tuning (міст розд. 1 і 2)", "tuned"],
        ],
        col_widths_cm=[3.5, 8, 3.5])

    _heading(doc, "Додаток Б. Приклади CLI-команд", level=2)
    _para(doc, "Швидкий прогон без MLflow і без tuning:")
    _code(doc,
        "python scripts/run_experiment.py --dataset phiusiil \\\n"
        "  --sample 5000 --folds 3 --no-mlflow --no-tuning"
    )
    _para(doc, "Повний прогон з усіма 12 моделями:")
    _code(doc,
        "python scripts/run_experiment.py --dataset steel_plate --folds 5"
    )
    _para(doc, "Тільки конкретні моделі:")
    _code(doc,
        "python scripts/run_experiment.py --dataset loan_approval \\\n"
        "  --sample 8000 --models gam,tuned_gam,cma_mixture"
    )

    _heading(doc, "Додаток В. Структура файлів програми", level=2)
    _image(doc, FIGURES_DIR / "01_architecture.png", width_cm=16,
           caption="Рисунок В.1 — Блочна схема модулів програми")


# ============================================================================
# збірка

def build_document() -> Path:
    doc = Document()
    # дефолтний шрифт для всього документа — Times New Roman 12pt
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    # сторінкові поля
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    _add_title_page(doc)
    _add_annotation(doc)
    _add_toc(doc)
    _add_intro(doc)
    _add_section_1(doc)
    _add_section_2(doc)
    _add_section_3(doc)
    _add_section_4(doc)
    _add_section_5(doc)
    _add_section_6(doc)
    _add_section_7(doc)
    _add_section_8(doc)
    _add_conclusions(doc)
    _add_appendices(doc)

    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    return OUT_PATH


def main() -> int:
    out = build_document()
    size_kb = out.stat().st_size // 1024
    print(f"Згенеровано: {out} ({size_kb} KB)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
