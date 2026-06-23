"""Генерує «Завдання на випускну кваліфікаційну роботу» для Богдана.

Структура і стиль повторюють зразок Пашняка Т. О.:
шапка ЗВО, ЗАТВЕРДЖУЮ Зав. кафедри, ПІБ студента, тема, вихідні дані,
перелік питань, перелік графічного матеріалу, консультанти (таблиця),
КАЛЕНДАРНИЙ ПЛАН (таблиця з 12 етапів), підписи.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt


REPO_ROOT = Path(__file__).resolve().parent.parent
OUT = REPO_ROOT.parent / "Волошенюк_Б_А_завдання_друкується_окремо.docx"


def _set_run(run, *, size=12, bold=False, italic=False,
             font="Times New Roman"):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rfonts.set(qn("w:cs"), font)


def _para(doc, text, *, size=12, bold=False, italic=False,
          align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=None,
          space=Pt(2)):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    if indent is not None:
        pf.first_line_indent = indent
    pf.space_after = space
    pf.line_spacing = 1.15
    run = p.add_run(text)
    _set_run(run, size=size, bold=bold, italic=italic)
    return p


def _caption(doc, text):
    """Маленька курсивна підказка (наприклад, «(прізвище, ім'я, по батькові)»)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_after = Pt(4)
    pf.line_spacing = 1.0
    run = p.add_run(text)
    _set_run(run, size=10, italic=True)


def _setup_page(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)


def _make_table(doc, headers, rows, *, widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        _set_run(run, size=12, bold=True)

    for r_i, row in enumerate(rows, start=1):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i].cells[c_i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = (WD_ALIGN_PARAGRAPH.LEFT if c_i == 1
                           else WD_ALIGN_PARAGRAPH.CENTER)
            run = p.add_run(str(val))
            _set_run(run, size=11)
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    return table


def build_doc() -> Document:
    doc = Document()
    _setup_page(doc)

    # ---- Шапка ЗВО ----
    _para(doc,
          "Чернівецький національний університет імені Юрія Федьковича",
          align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    _caption(doc, "(назва ЗВО)")

    _para(doc,
          "Інститут фізико-технічних та комп'ютерних наук",
          align=WD_ALIGN_PARAGRAPH.LEFT)
    _para(doc,
          "Кафедра математичних проблем управління і кібернетики",
          align=WD_ALIGN_PARAGRAPH.LEFT)
    _para(doc,
          "Спеціальність 122 – Комп'ютерні науки",
          align=WD_ALIGN_PARAGRAPH.LEFT,
          space=Pt(10))

    # ---- ЗАТВЕРДЖУЮ ----
    _para(doc, "ЗАТВЕРДЖУЮ", bold=True,
          align=WD_ALIGN_PARAGRAPH.RIGHT)
    _para(doc,
          "Зав. кафедри ________________ Ігор МАЛИК",
          align=WD_ALIGN_PARAGRAPH.RIGHT)
    _para(doc,
          "«26» серпня 2025 р.",
          align=WD_ALIGN_PARAGRAPH.RIGHT,
          space=Pt(12))

    # ---- Заголовок ----
    _para(doc,
          "ЗАВДАННЯ НА ВИПУСКНУ КВАЛІФІКАЦІЙНУ РОБОТУ",
          size=14, bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER,
          space=Pt(14))

    # ---- Студент ----
    _para(doc,
          "студенту  Волошенюку Богдану Анатолійовичу",
          align=WD_ALIGN_PARAGRAPH.LEFT)
    _caption(doc, "(прізвище, ім'я, по батькові)")

    # ---- 1. Тема ----
    _para(doc,
          "1. Тема: Порівняння методів класифікації з розширеним "
          "алгоритмом CMA-ES",
          bold=True)

    # ---- 2. Затверджена ----
    _para(doc,
          "2. Затверджена на засіданні кафедри МПУіК, протокол № 1 "
          "від «26» серпня 2025 р.")
    _para(doc,
          "Термін подачі студентом завершеної роботи на кафедру "
          "«05» червня 2026 р.")

    # ---- 3. Вихідні дані ----
    _para(doc,
          "3. Вихідні дані до випускної кваліфікаційної роботи: "
          "операційна система Windows 11; середовище розробки "
          "Visual Studio Code; мова програмування Python 3.14; "
          "бібліотеки scikit-learn, NumPy, pandas, cma, MLflow, "
          "Streamlit, pytest; набори даних з UCI Machine Learning "
          "Repository – PhiUSIIL Phishing URL (ID 967, 2024), "
          "Steel Plate Defects (ID 198), Default of Credit Card "
          "Clients (ID 350); фіксоване значення зерна генератора "
          "випадкових чисел random_state = 42.")

    # ---- 4. Перелік питань ----
    _para(doc, "4. Перелік питань, що їх належить розробити:")
    _para(doc,
          "1. Опанувати теоретичні основи узагальнених адитивних "
          "моделей (GAM): базисне представлення гладких функцій, "
          "B-сплайни, параметри згладжування та функції зв'язку.")
    _para(doc,
          "2. Опанувати алгоритм адаптації коваріаційної матриці "
          "CMA-ES і його розширення на суміш нормальних розподілів "
          "з EM-оновленням параметрів.")
    _para(doc,
          "3. Розробити програмний модуль classification_cma_es з "
          "єдиним sklearn-сумісним інтерфейсом для дванадцяти "
          "класифікаторів (5 базових, 2 CMA-NN, 5 tuned-моделей з "
          "автоматичним підбором гіперпараметрів через CMA-ES).")
    _para(doc,
          "4. Самостійно реалізувати GAM-класифікатор на основі "
          "B-сплайнового перетворення та розширений варіант CMA-ES "
          "зі сумішами нормальних розподілів.")
    _para(doc,
          "5. Сформувати методику оцінки моделей: стратифікований "
          "поділ train/test, Stratified K-Fold перехресна валідація, "
          "три класичні метрики класифікації (accuracy, F1-score, "
          "ROC-AUC) та провести експерименти на трьох сучасних "
          "відкритих датасетах.")

    # ---- 5. Графічний матеріал ----
    _para(doc,
          "5. Перелік графічного матеріалу (з точним зазначенням "
          "обов'язкових креслень): схема архітектури програмного "
          "модуля; зведені таблиці метрик 12 моделей на трьох "
          "датасетах; матриці плутанини та ROC-криві лідируючих "
          "моделей; криві збіжності алгоритму CMA-ES; графік "
          "порівняння F1-score 12 моделей на трьох датасетах; "
          "візуалізація коваріаційної матриці CMA-ES.")

    # ---- 6. Консультанти (таблиця) ----
    _para(doc,
          "6. Консультанти по роботі, із зазначенням розділів роботи, "
          "що стосуються їх", space=Pt(4))

    _make_table(doc,
        headers=["Розділ", "Консультант", "Завдання видав", "Завдання прийняв"],
        rows=[
            ["Розділ", "Консультант", "Підпис, дата", "Підпис, дата"],
        ],
        widths=[3, 5, 4, 4])
    doc.add_paragraph()

    # ---- 7. Дата видачі ----
    _para(doc,
          "7. Дата видачі завдання   ____________________________________"
          "_____________")

    _para(doc, "Керівник   _____________________________________________",
          space=Pt(0))
    _caption(doc, "(підпис)")

    _para(doc,
          "Завдання прийняв до виконання   _________________________",
          space=Pt(0))
    _caption(doc, "(підпис)")

    # ---- КАЛЕНДАРНИЙ ПЛАН ----
    _para(doc, "КАЛЕНДАРНИЙ ПЛАН",
          size=14, bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER,
          space=Pt(8))

    schedule_rows = [
        ["1", "Отримання завдання на випускну кваліфікаційну роботу",
         "01.09.25 р."],
        ["2", "Опрацювання рекомендованих джерел", "14.10.25 р."],
        ["3",
         "Опрацювання теоретичних основ узагальнених адитивних "
         "моделей (GAM)",
         "02.12.25 р."],
        ["4",
         "Опрацювання теорії алгоритму CMA-ES і його розширення "
         "сумішами розподілів",
         "16.12.25 р."],
        ["5",
         "Обґрунтування вибору датасетів та проектування "
         "архітектури програмного модуля",
         "31.12.25 р."],
        ["6",
         "Реалізація базових класифікаторів та GAM-класифікатора "
         "власної реалізації",
         "24.01.26 р."],
        ["7",
         "Реалізація розширеного CMA-ES зі сумішами розподілів та "
         "механізму підбору гіперпараметрів",
         "17.02.26 р."],
        ["8",
         "Проведення експериментів на трьох датасетах і обчислення "
         "метрик якості",
         "28.02.26 р."],
        ["9",
         "Зведений аналіз результатів та формулювання висновків "
         "дослідження",
         "13.03.26 р."],
        ["10", "Оформлення пояснювальної записки", "29.03.26 р."],
        ["11", "Представлення готової роботи", "05.06.26 р."],
        ["12",
         "Захист випускної кваліфікаційної роботи",
         "згідно з розкладом"],
    ]

    _make_table(doc,
        headers=["№ п/п", "Назва етапів випускної кваліфікаційної роботи",
                 "Термін виконання етапів"],
        rows=schedule_rows,
        widths=[1.5, 11.0, 4.0])

    doc.add_paragraph()

    # ---- Підписи ----
    _para(doc,
          "Студент\t\t__________________",
          space=Pt(0))
    _caption(doc, "(підпис)")

    _para(doc,
          "Керівник роботи\t__________________",
          space=Pt(0))
    _caption(doc, "(підпис)")

    return doc


def main() -> int:
    doc = build_doc()
    doc.save(str(OUT))
    print(f"OK: {OUT.name} ({OUT.stat().st_size // 1024} KB)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
