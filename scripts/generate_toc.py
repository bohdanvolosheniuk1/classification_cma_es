"""Генерує окремий файл ``зміст.docx`` зі змістом всієї дипломної роботи.

Цей файл Богдан вставить на початок свого диплома вручну
(скопіювати → вставити перед розділом 1).

Виклик::

    python scripts/generate_toc.py
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from generate_diploma_sections import (
    _set_run, _h, add_table_of_contents,
)


REPO_ROOT = Path(__file__).resolve().parent.parent
OUT_PATH = REPO_ROOT.parent / "зміст.docx"


def main() -> int:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    add_table_of_contents(doc)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    size_kb = OUT_PATH.stat().st_size // 1024
    print(f"Згенеровано: {OUT_PATH} ({size_kb} KB)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
