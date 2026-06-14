"""Замінює OMML-формули в розділах 1-2 на PNG з нумерацією.

OMML-формули, успадковані з оригінального .docx, погано рендеряться в
LibreOffice, Word на Mac та при PDF-конвертації. Натомість PNG-формули
працюють скрізь. Скрипт читає ``диплом 1_BACKUP_polished.docx``,
проходить по параграфам, замінює основні блочні формули на PNG із
нумерацією (1.1)–(1.9) та (2.1)–(2.5), а допоміжні короткі OMML-руни
у блоках «Де:» прибирає (вони і так дублюють опис змінних поряд).
Додатково вирівнює "Де:"-блоки на лівий край (зараз вони центровані).

Запуск::

    python scripts/patch_chapter12_formulas.py
"""

from __future__ import annotations

import re
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

REPO_ROOT = Path(__file__).resolve().parent.parent
SRC = REPO_ROOT.parent / "диплом 1_BACKUP_polished.docx"
DST = REPO_ROOT.parent / "диплом 1_BACKUP_final.docx"
FIGURES_DIR = REPO_ROOT.parent / "diploma_figures"


# Зіставлення позиції OMML-параграфа з PNG і номером формули.
# Позиції визначено через _dump-аналіз polished-файлу.
# Решта OMML-параграфів – це короткі inline-руни в "Де:"-блоках,
# їх просто видаляємо.
FORMULA_MAP = {
    12:  ("formula_1_1_additive.png", "(1.1)", 11.0),
    20:  ("formula_1_2_linear.png",   "(1.2)", 13.0),
    25:  ("formula_1_3_gam_link.png", "(1.3)", 12.0),
    37:  ("formula_1_4_basis.png",    "(1.4)", 10.0),
    46:  ("formula_1_5_poly.png",     "(1.5)", 13.0),
    74:  ("formula_1_6_loss.png",     "(1.6)", 11.0),
    91:  ("formula_1_3_gam_link.png", "(1.3)", 12.0),  # повторне посилання
    99:  ("formula_1_7_identity.png", "(1.7)",  6.0),
    105: ("formula_1_8_log.png",      "(1.8)",  7.0),
    112: ("formula_1_9_logit.png",    "(1.9)",  8.0),
    156: ("formula_2_1_sample.png",   "(2.1)", 12.0),
    170: ("formula_2_2_sample2.png",  "(2.2)", 11.0),
    180: ("formula_2_3_mean.png",     "(2.3)", 11.0),
    190: ("formula_2_1_sample.png",   "(2.1)", 12.0),  # повторне посилання
    200: ("formula_2_3_mean.png",     "(2.3)", 11.0),  # повторне посилання
    215: ("formula_2_4_cov.png",      "(2.4)", 15.0),
    234: ("formula_2_5_sigma.png",    "(2.5)", 14.0),
}


def _para_has_block_formula(p) -> bool:
    return "<m:oMathPara" in p._p.xml


def _para_has_inline_math(p) -> bool:
    return "<m:oMath" in p._p.xml and "<m:oMathPara" not in p._p.xml


def _clear_paragraph(p) -> None:
    """Видаляє всі дочірні елементи параграфа, окрім pPr."""
    p_elem = p._p
    pPr = p_elem.find(qn("w:pPr"))
    for child in list(p_elem):
        if child.tag != qn("w:pPr"):
            p_elem.remove(child)


M_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"
W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


# Mathematical italic Unicode → звичайна латиниця/грецька.
# Word у Cambria Math рендерить ці гліфи як formula italic, але звичайний
# Times New Roman їх часто не має. Робимо ASCII-конверсію.
_MATH_ITALIC_MAP = {}
# Latin A-Z (Math Italic): U+1D434..U+1D44D, A=U+1D434
for i, ch in enumerate("ABCDEFGHIJKLMNOPQRSTUVWXYZ"):
    _MATH_ITALIC_MAP[chr(0x1D434 + i)] = ch
# Latin a-z: U+1D44E..U+1D467, a=U+1D44E, h missing → U+210E
for i, ch in enumerate("abcdefghijklmnopqrstuvwxyz"):
    if ch == 'h':
        _MATH_ITALIC_MAP[chr(0x210E)] = 'h'
    else:
        _MATH_ITALIC_MAP[chr(0x1D44E + i)] = ch
# Грецькі italic Math: U+1D6FC..U+1D71B (α..ω) – і ψ/ω
_greek_italic_start = 0x1D6FC  # α
_greek_lower = "αβγδεζηθικλμνξοπρςστυφχψω"
for i, ch in enumerate(_greek_lower):
    _MATH_ITALIC_MAP[chr(_greek_italic_start + i)] = ch
# Math bold Greek upper U+1D6A8..U+1D6C2
_greek_upper = "ΑΒΓΔΕΖΗΘΙΚΛΜΝΞΟΠΡϴΣΤΥΦΧΨΩ"
_GREEK_ITAL_UPPER_START = 0x1D6E2  # Math italic Greek Cap A
for i, ch in enumerate(_greek_upper):
    _MATH_ITALIC_MAP[chr(_GREEK_ITAL_UPPER_START + i)] = ch
# Math bold italic alternatives – поки що не покриваємо


def _to_plain_text(s: str) -> str:
    """Конвертує math-italic Unicode у звичайний текст."""
    return "".join(_MATH_ITALIC_MAP.get(c, c) for c in s)


_SUPSCRIPT_DIGITS_MAP = {
    "0": "⁰", "1": "¹", "2": "²", "3": "³", "4": "⁴",
    "5": "⁵", "6": "⁶", "7": "⁷", "8": "⁸", "9": "⁹",
    "+": "⁺", "-": "⁻", "=": "⁼", "(": "⁽", ")": "⁾",
    "n": "ⁿ", "i": "ⁱ", "T": "ᵀ", "t": "ᵗ",
}


def _to_unicode_superscript(s: str) -> str:
    """Перетворює рядок на Unicode superscript-форму, де можливо."""
    return "".join(_SUPSCRIPT_DIGITS_MAP.get(ch, ch) for ch in s)


_SUBSCRIPT_DIGITS = {ch: chr(0x2080 + i) for i, ch in enumerate("0123456789")}
_SUBSCRIPT_DIGITS["+"] = "₊"
_SUBSCRIPT_DIGITS["-"] = "₋"
_SUBSCRIPT_DIGITS["="] = "₌"
_SUBSCRIPT_DIGITS["("] = "₍"
_SUBSCRIPT_DIGITS[")"] = "₎"
_SUBSCRIPT_LETTERS = {
    "a": "ₐ", "e": "ₑ", "h": "ₕ", "i": "ᵢ", "j": "ⱼ", "k": "ₖ",
    "l": "ₗ", "m": "ₘ", "n": "ₙ", "o": "ₒ", "p": "ₚ", "r": "ᵣ",
    "s": "ₛ", "t": "ₜ", "u": "ᵤ", "v": "ᵥ", "x": "ₓ",
}


def _to_unicode_subscript(s: str) -> str:
    """Перетворює рядок на Unicode subscript-форму (для індексів формул)."""
    out = []
    for ch in s:
        if ch in _SUBSCRIPT_DIGITS:
            out.append(_SUBSCRIPT_DIGITS[ch])
        elif ch.lower() in _SUBSCRIPT_LETTERS:
            out.append(_SUBSCRIPT_LETTERS[ch.lower()])
        else:
            out.append(ch)
    return "".join(out)


def _omml_to_plain(elem) -> str:
    """Розкручує OMML-структуру у простий читабельний текст.

    Підтримує: <m:t> – звичайний текст, <m:d> – дужки (fenced),
    <m:sub> – підіндекс (Unicode subscript), <m:sup> – надіндекс
    (^xxx нотація). Працює рекурсивно.
    """
    parts = []
    for child in elem:
        local = child.tag.split("}", 1)[-1]
        if local == "t":
            parts.append(child.text or "")
        elif local == "d":
            # дужки. Тип дужок з dPr/begChr/endChr; за замовчуванням ( )
            beg, end = "(", ")"
            pr = child.find(qn("m:dPr"))
            if pr is not None:
                bch = pr.find(qn("m:begChr"))
                if bch is not None and bch.get(qn("m:val")):
                    beg = bch.get(qn("m:val"))
                ech = pr.find(qn("m:endChr"))
                if ech is not None and ech.get(qn("m:val")):
                    end = ech.get(qn("m:val"))
            inner = _omml_to_plain(child)
            parts.append(f"{beg}{inner}{end}")
        elif local == "sub":
            # піднідекс – у Unicode subscript
            inner = _omml_to_plain(child)
            parts.append(_to_unicode_subscript(inner))
        elif local == "sup":
            inner = _omml_to_plain(child)
            if not inner:
                continue
            # пробуємо Unicode superscript; якщо не всі символи перевелись –
            # ставимо ^ перед оригіналом
            sup = _to_unicode_superscript(inner)
            if all(ch in _SUPSCRIPT_DIGITS_MAP for ch in inner):
                parts.append(sup)
            else:
                parts.append(f"^{inner}")
        else:
            # рекурсивно пройти всередину
            parts.append(_omml_to_plain(child))
    return "".join(parts)


def _collect_omml_text(elem) -> str:
    """Збирає текст з OMML з урахуванням дужок та підіндексів."""
    return _omml_to_plain(elem)


def _make_text_run(text: str, *, italic: bool = True, size_pt: int = 14):
    """Створює <w:r>-елемент із заданим текстом (italic + Times New Roman)."""
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    rPr.append(rFonts)
    if italic:
        i_elem = OxmlElement("w:i")
        rPr.append(i_elem)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size_pt * 2))
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _strip_omml_keep_text(p) -> bool:
    """Замість OMML-елементів вставляє звичайний w:r-текст з курсивом.

    Тобто символи `𝛾`, `𝛽₀`, `𝑓ᵢ𝑥ᵢ` (math-italic Unicode, що погано
    рендериться) перетворюються на звичайні `γ`, `β₀`, `fᵢxᵢ` у курсиві.

    Повертає True, якщо була принаймні одна модифікація.
    """
    p_elem = p._p
    changed = False
    for child in list(p_elem):
        tag = child.tag
        if tag.startswith(M_NS):
            # Витягуємо текст з OMML
            text = _collect_omml_text(child)
            plain = _to_plain_text(text)
            if plain.strip():
                # Створюємо звичайний run-текст і вставляємо ПЕРЕД OMML
                new_r = _make_text_run(plain, italic=True, size_pt=14)
                p_elem.insert(list(p_elem).index(child), new_r)
            # Видаляємо OMML
            p_elem.remove(child)
            changed = True
    return changed


def _insert_png_with_number(p, png_path: Path, number: str,
                            width_cm: float) -> None:
    """Перетворює параграф на: [PNG по центру] [tab] [номер праворуч]."""
    _clear_paragraph(p)
    pf = p.paragraph_format
    pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0)
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    # tab stop праворуч для номера
    pf.tab_stops.add_tab_stop(Cm(16.0), WD_PARAGRAPH_ALIGNMENT.RIGHT)
    # картинка
    run_pic = p.add_run()
    run_pic.add_picture(str(png_path), width=Cm(width_cm))
    # таб + номер
    run_num = p.add_run(f"\t{number}")
    run_num.font.name = "Times New Roman"
    run_num.font.size = Pt(14)


def _is_where_marker(text: str) -> bool:
    """Чи це параграф з 'Де:' маркером."""
    t = text.strip().lower()
    return t in ("де:", "де :")


# Випадок 1: "симв - опис" або "симв – опис" (з пробілом або без)
_DEF_RE_FULL = re.compile(r"^\s*(\S.*?)\s*[-––]\s*(.+?)\s*$")
# Випадок 2: "- опис" (без символу – це продовження попереднього)
_DEF_RE_NOSYM = re.compile(r"^\s*[-––]\s*(.+?)\s*$")


def _looks_like_definition(text: str) -> bool:
    """Параграф виглядає як визначення (з символом або без)."""
    t = text.strip()
    if not t:
        return False
    return bool(_DEF_RE_NOSYM.match(t) or _DEF_RE_FULL.match(t))


def _bulletize_definition(p) -> bool:
    """Перетворює визначення на буллет '– симв – опис' або '– опис'.

    Поточний текст парситься на (symbol?, description). Параграф
    очищується і збирається з буллетом, italic-символом (якщо є) та
    звичайним описом. Уніфікує стиль із розд. 3-4.
    """
    text = p.text.strip().replace("\xa0", " ")
    text = " ".join(text.split())  # схлопуємо множинні пробіли

    symbol = ""
    descr = ""
    # Спочатку перевіряємо паттерн "- опис" (без символу)
    if text.lstrip().startswith(("-", "–", "–")):
        m = _DEF_RE_NOSYM.match(text)
        if m:
            descr = m.group(1).strip()
    else:
        # Інакше шукаємо повний паттерн "симв - опис"
        m = _DEF_RE_FULL.match(text)
        if m:
            symbol = m.group(1).strip()
            descr = m.group(2).strip()

    if not descr:
        return False

    _clear_paragraph(p)
    pf = p.paragraph_format
    pf.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    pf.left_indent = Cm(0.75)
    pf.first_line_indent = Cm(-0.5)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.4

    # буллет
    r1 = p.add_run("–  ")
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(14)
    if symbol:
        # символ у курсиві
        r2 = p.add_run(symbol)
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(14)
        r2.italic = True
        # тире і опис
        r3 = p.add_run(f" – {descr}")
    else:
        # тільки опис
        r3 = p.add_run(descr)
    r3.font.name = "Times New Roman"
    r3.font.size = Pt(14)
    return True


def _normalize_de_marker(p) -> bool:
    """Нормалізує текст параграфа 'Де:' / 'де:' до жирного 'Де:'.

    Повертає True, якщо параграф був маркером (для подальшого dedup).
    """
    t = p.text.strip().lower()
    if t not in ("де:", "де :"):
        return False
    # Замість усього вмісту – єдиний жирний run "Де:"
    _clear_paragraph(p)
    pf = p.paragraph_format
    pf.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0)
    pf.space_before = Pt(4)
    pf.space_after = Pt(2)
    run = p.add_run("Де:")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    return True


def _delete_paragraph(p) -> None:
    """Повністю видаляє параграф з body."""
    p_elem = p._p
    p_elem.getparent().remove(p_elem)


def _insert_chapter_heading_before(p, title: str, *, page_break: bool = True):
    """Вставляє ЗАГОЛОВОК 'РОЗДІЛ N. НАЗВА' (жирним, по центру) перед p.

    Опційно – з page-break перед заголовком, щоб розділ починався з
    нової сторінки.
    """
    p_elem = p._p
    parent = p_elem.getparent()
    idx = list(parent).index(p_elem)

    # 1. Параграф із page-break (порожній, тільки <w:br type="page"/>)
    if page_break:
        pb = OxmlElement("w:p")
        r = OxmlElement("w:r")
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        r.append(br)
        pb.append(r)
        parent.insert(idx, pb)
        idx += 1

    # 2. Параграф із заголовком розділу
    p_head = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    pPr.append(jc)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "240")
    spacing.set(qn("w:after"), "240")
    pPr.append(spacing)
    p_head.append(pPr)

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    b = OxmlElement("w:b")
    rPr.append(b)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "28")  # 14pt = 28 half-points
    rPr.append(sz)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    rPr.append(rFonts)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = title
    r.append(t)
    p_head.append(r)
    parent.insert(idx, p_head)


def _insert_page_break_before(p):
    """Вставляє page-break безпосередньо перед параграфом p."""
    p_elem = p._p
    parent = p_elem.getparent()
    idx = list(parent).index(p_elem)
    pb = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    pb.append(r)
    parent.insert(idx, pb)


def _make_subheading(p):
    """Перетворює параграф на жирний підпункт по центру (1.1, 2.3 і т.д.)."""
    text = p.text.strip()
    if not text:
        return
    _clear_paragraph(p)
    pf = p.paragraph_format
    pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0)
    pf.space_before = Pt(12)
    pf.space_after = Pt(8)
    pf.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


_SUBHEADING_RE = re.compile(r"^\d+\.\d+(\.\d+)?\s+\S")
_CONCLUSION_RE = re.compile(r"^Висновки до розділу\s+\d+\s*$", re.IGNORECASE)


def main() -> int:
    if not SRC.exists():
        print(f"ПОМИЛКА: {SRC} не знайдено", file=sys.stderr)
        return 1

    doc = Document(SRC)
    paragraphs = list(doc.paragraphs)

    n_replaced = 0
    n_stripped = 0
    n_aligned = 0
    n_de_normalized = 0

    for i, p in enumerate(paragraphs):
        text = p.text.strip()

        # 1. Замінити цільові формули на PNG з нумерацією
        if i in FORMULA_MAP:
            png_name, number, width_cm = FORMULA_MAP[i]
            png_path = FIGURES_DIR / png_name
            if not png_path.exists():
                print(f"  ПОПЕРЕДЖЕННЯ: {png_path} не знайдено")
                continue
            _insert_png_with_number(p, png_path, number, width_cm)
            n_replaced += 1
            continue

        # 2. У решті OMML-параграфів видаляти ЛИШЕ OMML-елементи,
        # текстові w:r-руни ("- залежна змінна," тощо) лишаються.
        if "<m:oMath" in p._p.xml:
            if _strip_omml_keep_text(p):
                n_stripped += 1

        # 3. Нормалізувати "Де:"-маркери (жирний + лівий край)
        if _normalize_de_marker(p):
            n_de_normalized += 1
            continue

        # 4. Перебудувати рядки-визначення у формат "– симв – опис"
        if _looks_like_definition(p.text):
            if _bulletize_definition(p):
                n_aligned += 1

    # 4.5. Глобальна заміна em-dash "—" на en-dash "–" у всьому тексті
    # (загальна вимога куратора).
    n_dash_replaced = 0
    for p in doc.paragraphs:
        for r in p.runs:
            if r.text and ("—" in r.text or "•" in r.text):
                new_text = r.text.replace("—", "–").replace("•", "–")
                if new_text != r.text:
                    r.text = new_text
                    n_dash_replaced += 1

    # 5. Видалити дублікати "Де:" – якщо два маркери підряд, залишаємо
    # тільки перший.
    paragraphs = list(doc.paragraphs)
    n_duplicates = 0
    for i in range(len(paragraphs) - 1, 0, -1):
        cur = paragraphs[i].text.strip().lower()
        prev = paragraphs[i - 1].text.strip().lower()
        if cur in ("де:", "де :") and prev in ("де:", "де :"):
            _delete_paragraph(paragraphs[i])
            n_duplicates += 1

    # 6. Зробити жирний-центр підпункти "1.1", "1.2", "2.1" і т.д.
    # Висновки до розділу – теж жирний центр + page-break before.
    paragraphs = list(doc.paragraphs)
    n_subheadings = 0
    n_conclusion_breaks = 0
    for p in paragraphs:
        t = p.text.strip()
        if _SUBHEADING_RE.match(t):
            _make_subheading(p)
            n_subheadings += 1
        elif _CONCLUSION_RE.match(t):
            _make_subheading(p)
            _insert_page_break_before(p)
            n_conclusion_breaks += 1

    # 7. Перед першим підпунктом '1.1' додаємо заголовок РОЗДІЛ 1,
    # перед '2.1' – РОЗДІЛ 2 з page-break.
    paragraphs = list(doc.paragraphs)
    n_chapter_headings = 0
    seen_chapter1 = False
    seen_chapter2 = False
    for p in paragraphs:
        t = p.text.strip()
        if not seen_chapter1 and t.startswith("1.1 "):
            _insert_chapter_heading_before(
                p, "РОЗДІЛ 1. УЗАГАЛЬНЕНІ АДИТИВНІ МОДЕЛІ",
                page_break=True)
            seen_chapter1 = True
            n_chapter_headings += 1
        elif not seen_chapter2 and t.startswith("2.1 "):
            _insert_chapter_heading_before(
                p, "РОЗДІЛ 2. АЛГОРИТМ АДАПТАЦІЇ КОВАРІАЦІЙНОЇ МАТРИЦІ",
                page_break=True)
            seen_chapter2 = True
            n_chapter_headings += 1

    doc.save(DST)
    print(f"Замінено формул на PNG     : {n_replaced}")
    print(f"OMML-руни прибрано (текст залишено): {n_stripped}")
    print(f"'Де:' нормалізовано (жирний): {n_de_normalized}")
    print(f"Дублікатів 'Де:' видалено   : {n_duplicates}")
    print(f"Визначень буллетизовано   : {n_aligned}")
    print(f"Підпунктів-заголовків      : {n_subheadings}")
    print(f"Висновків з page-break     : {n_conclusion_breaks}")
    print(f"Заголовків розділів додано : {n_chapter_headings}")
    print(f"Збережено: {DST.name} ({DST.stat().st_size // 1024} KB)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
