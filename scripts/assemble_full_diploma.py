"""Збирає повну дипломну роботу у один файл диплом 1.docx.

Структура фінального документа:

1. Зміст (TOC)
2. Оригінальні розділи 1 і 2 (теоретична частина Богдана)
3. Розділ 3 — програмна реалізація
4. Розділ 4 — результати експериментів (з рисунками)
5. Список використаних джерел

Перед запуском має існувати ``диплом 1_BACKUP_до_інтеграції.docx``
з оригіналом. Якщо його немає, скрипт зробить його з поточного
диплома (тільки якщо у тому немає вже доданих розділів — інакше
відмовиться щоб не зіпсувати оригінал).

Запуск::

    python scripts/assemble_full_diploma.py
"""

from __future__ import annotations

import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from docxcompose.composer import Composer

from generate_diploma_sections import (
    add_annotation,
    add_appendix,
    add_bibliography,
    add_conclusions,
    add_introduction,
    add_sections_to,
    add_table_of_contents,
    add_title_page,
)


REPO_ROOT = Path(__file__).resolve().parent.parent
DIPLOMA_PATH = REPO_ROOT.parent / "випускна кваліфікаційна робота.docx"
BACKUP_PATH = REPO_ROOT.parent / "диплом 1_BACKUP_до_інтеграції.docx"
MERGED_PATH = REPO_ROOT.parent / "диплом 1_BACKUP_merged.docx"
POLISHED_PATH = REPO_ROOT.parent / "диплом 1_BACKUP_polished.docx"
FINAL_BACKUP_PATH = REPO_ROOT.parent / "диплом 1_BACKUP_final.docx"


def _make_blank_doc() -> Document:
    """Створює пустий Document зі стандартними налаштуваннями сторінки."""
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
    return doc


def main() -> int:
    if not BACKUP_PATH.exists():
        print(f"ПОМИЛКА: не знайдено бекап {BACKUP_PATH}",
              file=sys.stderr)
        print("Спочатку запустіть integrate_sections_to_diploma.py хоч раз "
              "щоб створити бекап.", file=sys.stderr)
        return 1

    with tempfile.TemporaryDirectory() as tmpdir:
        tmp = Path(tmpdir)

        # 1. Титулка + Анотація + Зміст + Вступ — головна частина перед теорією
        print("Готую титулку, анотацію, зміст і вступ...")
        front_doc = _make_blank_doc()
        add_title_page(front_doc)
        add_annotation(front_doc)
        add_table_of_contents(front_doc)
        front_doc.add_page_break()
        add_introduction(front_doc)
        front_path = tmp / "01_front.docx"
        front_doc.save(str(front_path))

        # 2. Оригінал теоретичної частини (розділи 1 і 2 Богдана).
        # Пріоритет: final (з PNG-формулами) > polished > merged > оригінал.
        if FINAL_BACKUP_PATH.exists():
            source = FINAL_BACKUP_PATH
        elif POLISHED_PATH.exists():
            source = POLISHED_PATH
        elif MERGED_PATH.exists():
            source = MERGED_PATH
        else:
            source = BACKUP_PATH
        print(f"Беру теоретичну частину з: {source.name}")
        original_path = tmp / "02_original.docx"
        original_path.write_bytes(source.read_bytes())

        # 3. Розділи 3, 4 + загальні висновки + бібліографія + додаток
        print("Готую розділи 3, 4, висновки, список літератури і додаток...")
        sections_doc = _make_blank_doc()
        sections_doc.add_paragraph()
        add_sections_to(sections_doc)
        add_conclusions(sections_doc)
        add_bibliography(sections_doc)
        add_appendix(sections_doc)
        sections_path = tmp / "03_sections.docx"
        sections_doc.save(str(sections_path))

        # 4. Зливаємо все разом через docxcompose
        print("Зливаю всі частини в один файл...")
        master = Document(str(front_path))
        composer = Composer(master)
        composer.append(Document(str(original_path)))
        composer.append(Document(str(sections_path)))

        master.save(str(DIPLOMA_PATH))

    size_kb = DIPLOMA_PATH.stat().st_size // 1024
    print(f"\nГотово: {DIPLOMA_PATH.name} ({size_kb} KB)")

    # коротка перевірка вмісту
    doc = Document(str(DIPLOMA_PATH))
    total = len(doc.paragraphs)
    tables = len(doc.tables)
    inlines = sum(1 for _ in doc.inline_shapes)
    print(f"  Параграфів: {total}")
    print(f"  Таблиць: {tables}")
    print(f"  Зображень: {inlines}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
